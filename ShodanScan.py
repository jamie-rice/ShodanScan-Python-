'''ShodanSearch
    # -  a script to search shodan for vulnerable devices by Jamie Rice - 2024 
    # - use case is a client requesting security researcher to find all devices which are vulnerable to a specific CVE
    # - Perhaps they wish to find clients who have not updated a system from the vulnerable version? 
    # - Well the researcher can use this to search, for example, a good one is CVE-2024-24989 
    # - CVE-2024-24989, product:nginx version1.25.3

Importing required libraries for the following script
    # - Shodan for the main function of the script, 
    # - docx to produce a word document which formates the output
    # - datetime to print the time of execution on the doc
'''

import shodan 
from docx import Document
from datetime import datetime

#API key obtained from Shodan login dashboard
#Intilisaing shodan API object, allowing the use of the Shodan library 
api_key = 'A4nUAZNda1GRGU0dXApja5o3DhFZLuel'
api = shodan.Shodan(api_key)

#initialising datetime object to provide date//time on outputted document containing results
execution_time = datetime.now()
date_and_time_str = execution_time.strftime("%d/%m/%Y %H:%M:%S") # dd/mm/YY H:M:S

#encloses the shodan search and handles the API errors if it fails execution  
try: 
    #function peforming the foundational search of shodan
    # Takes the shodan api object and searches using provided 'query' (user input)
    def search_shodan(query):
        search_results = api.search(query)
        return search_results

    #The following block formats a series of print statements to screen 
    # - to tell user how to format input (particularly if more than one query)
    # - prompt the user to input this query
    print()
    print('Searchterms used are those used on shodan, for exmpale if we wish to type port & technology it would be ===> port:443 product:nginx')
    print()
    query=input('Enter your searchterms here: ')
    
    #PERFORMS THE SEARCH! uses user query that was inputed 
    results = search_shodan(query)

    #First line creates document called results_from_search
    #The rest of the lines set headings for the word document that results are being saved too
    results_from_search = Document()
    results_from_search.add_heading('ShodanScan: Uncovering Vulnerable Devices with Python', level=1)
    results_from_search.add_heading('Script by Jamie Rice, 2024', level=2)
    results_from_search.add_heading('Your Entered Search Query Was == ' + query, level=4)
    results_from_search.add_heading('Your search was produced at '+ date_and_time_str, level=4) #uses the variable define on line 17
    results_from_search.add_paragraph()
    results_from_search.add_paragraph("--------------------------")
    
    #initialises for loop to cycle through matched results from shodan search
    #extracts the information(the specified filters) and displays in word document
    for matched_result in results['matches']:
    
        #Intiliases placeholder with label, i.e IP: {} awaits recieved input, format marks what will go in {}
        #matched_result.get('ip_str') uses the key 'ip_str' to find associated value in matched_result dictionary and return shodan search value
        #some reuslts have N/A as they could potentially not have results - where as the other fields typically will when querying shodan
        results_from_search.add_paragraph('IP: {}'.format(matched_result.get('ip_str')))
        results_from_search.add_paragraph('Port: {}'.format(matched_result.get('port')))
        results_from_search.add_paragraph('Version: {}'.format(matched_result.get('version', 'N/A')))
        results_from_search.add_paragraph('Software: {}'.format(matched_result.get('product')))
        results_from_search.add_paragraph('Device: {}'.format(matched_result.get('devicetype')))
        results_from_search.add_paragraph('Organisation: {}'.format(matched_result.get('org', 'N/A')))
        results_from_search.add_paragraph('Location: {}'.format(matched_result.get('location', 'N/A')))
        results_from_search.add_paragraph()
        results_from_search.add_paragraph("--------------------------")
        results_from_search.add_paragraph()
    
    #saves document and uses the user inputted query as the text file 
    results_from_search.save((query) + ' Search results document.docx')

#part of the error handling process, will print errors regarding API execution
except shodan.APIError as e:
    print ('Error: %s' % e)


